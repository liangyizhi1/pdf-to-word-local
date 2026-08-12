from __future__ import annotations

import io
import os
import re
from collections.abc import Callable
from copy import deepcopy
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any

try:
    import pymupdf as fitz
except ImportError:
    try:
        import fitz
    except ImportError:
        fitz = None


class FormulaDependencyError(RuntimeError):
    """Raised when an explicitly requested formula component is unavailable."""


@dataclass(frozen=True)
class FormulaCandidate:
    page: int
    index: int
    bbox: tuple[float, float, float, float]
    width: int
    height: int
    image: bytes
    extension: str


@dataclass
class FormulaResult:
    page: int
    index: int
    bbox: list[float]
    status: str
    latex: str | None = None
    confidence: float | None = None
    rendering: str | None = None
    reason: str | None = None


@dataclass
class FormulaSummary:
    status: str
    engine: str = "rapid-latex-ocr"
    candidate_count: int = 0
    recognized_count: int = 0
    rejected_count: int = 0
    failed_count: int = 0
    native_equation_count: int = 0
    latex_fallback_count: int = 0
    appendix_added: bool = False
    results: list[FormulaResult] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


Recognizer = Callable[[bytes], Any]

_MATH_MARKERS = (
    r"\frac",
    r"\sqrt",
    r"\sum",
    r"\prod",
    r"\int",
    r"\lim",
    r"\begin",
    r"\alpha",
    r"\beta",
    r"\gamma",
    r"\theta",
    r"\lambda",
    r"\cdot",
    r"\times",
    "=",
    "^",
    "_",
    "≤",
    "≥",
)


def is_formula_image_block(
    block: dict[str, Any],
    *,
    page_width: float,
    page_height: float,
) -> bool:
    if block.get("type") != 1 or not isinstance(block.get("image"), (bytes, bytearray)):
        return False
    bbox = block.get("bbox")
    if not isinstance(bbox, (list, tuple)) or len(bbox) != 4:
        return False
    x0, y0, x1, y1 = (float(value) for value in bbox)
    width = max(0.0, x1 - x0)
    height = max(0.0, y1 - y0)
    if width < 24 or height < 12:
        return False
    page_area = max(1.0, page_width * page_height)
    if (width * height) / page_area > 0.35:
        return False
    if height > page_height * 0.25 or width > page_width * 0.96:
        return False
    aspect_ratio = width / max(height, 1.0)
    return 0.25 <= aspect_ratio <= 18.0


def extract_formula_candidates(
    source: str | Path,
    *,
    start_page: int = 1,
    end_page: int | None = None,
    max_per_page: int = 20,
) -> list[FormulaCandidate]:
    if fitz is None:
        raise FormulaDependencyError("Formula extraction requires PyMuPDF.")
    source_path = Path(source).expanduser().resolve()
    document = fitz.open(source_path)
    try:
        final_page = document.page_count if end_page is None else min(end_page, document.page_count)
        candidates: list[FormulaCandidate] = []
        for page_number in range(start_page - 1, final_page):
            page = document.load_page(page_number)
            blocks = page.get_text("dict").get("blocks", [])
            image_blocks = [
                block
                for block in blocks
                if is_formula_image_block(
                    block,
                    page_width=page.rect.width,
                    page_height=page.rect.height,
                )
            ]
            image_blocks.sort(key=lambda item: (item["bbox"][1], item["bbox"][0]))
            for index, block in enumerate(image_blocks[:max_per_page], start=1):
                candidates.append(
                    FormulaCandidate(
                        page=page_number + 1,
                        index=index,
                        bbox=tuple(round(float(value), 3) for value in block["bbox"]),
                        width=int(block.get("width") or 0),
                        height=int(block.get("height") or 0),
                        image=bytes(block["image"]),
                        extension=str(block.get("ext") or "png"),
                    )
                )
        return candidates
    finally:
        document.close()


def _prepare_image(image: bytes) -> bytes:
    try:
        from PIL import Image
    except ImportError:
        return image
    try:
        source = Image.open(io.BytesIO(image)).convert("RGBA")
        background = Image.new("RGBA", source.size, "white")
        background.alpha_composite(source)
        output = io.BytesIO()
        background.convert("RGB").save(output, format="PNG")
        return output.getvalue()
    except (OSError, ValueError):
        return image


def _load_recognizer() -> Recognizer:
    try:
        from rapid_latex_ocr import LatexOCR
    except ImportError as exc:
        raise FormulaDependencyError(
            "Formula OCR is not installed. Run install_formula_ocr.bat or install .[formula]."
        ) from exc
    try:
        engine = LatexOCR()
    except Exception as exc:
        raise FormulaDependencyError(f"Formula OCR could not start: {exc}") from exc
    return engine


def normalize_ocr_result(raw: Any) -> tuple[str, float | None]:
    confidence: float | None = None
    value = raw
    if isinstance(raw, tuple):
        value = raw[0] if raw else ""
    if isinstance(value, dict):
        confidence_value = value.get("confidence", value.get("score"))
        if isinstance(confidence_value, (int, float)):
            confidence = max(0.0, min(1.0, float(confidence_value)))
        value = value.get("latex", value.get("text", value.get("result", "")))
    latex = str(value or "").strip()
    if latex.startswith("$$") and latex.endswith("$$") and len(latex) >= 4:
        latex = latex[2:-2].strip()
    elif latex.startswith("$") and latex.endswith("$") and len(latex) >= 2:
        latex = latex[1:-1].strip()
    if latex.startswith(r"\[") and latex.endswith(r"\]"):
        latex = latex[2:-2].strip()
    return latex, confidence


def looks_like_formula(latex: str) -> bool:
    if len(latex) < 3 or len(latex) > 2000:
        return False
    if any(marker in latex for marker in _MATH_MARKERS):
        return True
    return bool(re.search(r"[A-Za-z0-9)]\s*[+\-*/<>]\s*[(A-Za-z0-9]", latex))


def _find_mml2omml_stylesheet() -> Path | None:
    configured = os.environ.get("MML2OMML_XSL")
    candidates = [Path(configured)] if configured else []
    for variable in ("ProgramFiles", "ProgramFiles(x86)"):
        root = os.environ.get(variable)
        if root:
            candidates.append(
                Path(root) / "Microsoft Office" / "root" / "Office16" / "MML2OMML.XSL"
            )
    return next((path for path in candidates if path.is_file()), None)


def _latex_to_omml(latex: str, stylesheet: Path) -> Any:
    try:
        from latex2mathml.converter import convert
        from lxml import etree
    except ImportError as exc:
        raise FormulaDependencyError(
            "Native Word equations require latex2mathml and lxml."
        ) from exc
    mathml = convert(latex)
    transform = etree.XSLT(etree.parse(str(stylesheet)))
    result = transform(etree.fromstring(mathml.encode("utf-8")))
    return deepcopy(result.getroot())


def _append_formulae(docx_path: Path, summary: FormulaSummary) -> None:
    recognized = [result for result in summary.results if result.status == "recognized"]
    if not recognized:
        return
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:
        raise FormulaDependencyError("Writing formulae to Word requires python-docx.") from exc

    document = Document(docx_path)
    document.add_page_break()
    document.add_heading("Recognized equations (experimental)", level=1)
    document.add_paragraph(
        "These equations were recognized from PDF image regions. Verify each equation against "
        "the source PDF before reuse."
    )
    stylesheet = _find_mml2omml_stylesheet()
    native_unavailable_reported = False
    for result in recognized:
        label = document.add_paragraph()
        label.add_run(f"Page {result.page}, formula {result.index}").bold = True
        equation_paragraph = document.add_paragraph()
        if stylesheet is not None:
            try:
                equation_paragraph._p.append(_latex_to_omml(result.latex or "", stylesheet))
                result.rendering = "native_word_equation"
                summary.native_equation_count += 1
                continue
            except Exception as exc:  # noqa: BLE001 - third-party boundary.
                if not native_unavailable_reported:
                    summary.warnings.append(
                        f"Native Word equation conversion was unavailable; LaTeX text was used: {exc}"
                    )
                    native_unavailable_reported = True
        run = equation_paragraph.add_run(result.latex or "")
        run.font.name = "Cambria Math"
        run.font.size = Pt(11)
        result.rendering = "latex_text"
        summary.latex_fallback_count += 1
    document.save(docx_path)
    summary.appendix_added = True


def enrich_docx_with_formulae(
    source: str | Path,
    docx_path: str | Path,
    *,
    start_page: int = 1,
    end_page: int | None = None,
    max_per_page: int = 20,
    recognizer: Recognizer | None = None,
) -> FormulaSummary:
    candidates = extract_formula_candidates(
        source,
        start_page=start_page,
        end_page=end_page,
        max_per_page=max_per_page,
    )
    if not candidates:
        return FormulaSummary(
            status="no_candidates",
            warnings=[
                (
                    "No image-based formula candidates were found. Version 0.2 does not yet "
                    "recognize formulae drawn as PDF text or vector paths."
                )
            ],
        )
    try:
        engine = recognizer or _load_recognizer()
    except FormulaDependencyError as exc:
        return FormulaSummary(
            status="unavailable",
            candidate_count=len(candidates),
            warnings=[str(exc)],
        )

    summary = FormulaSummary(status="completed", candidate_count=len(candidates))
    for candidate in candidates:
        result = FormulaResult(
            page=candidate.page,
            index=candidate.index,
            bbox=list(candidate.bbox),
            status="failed",
        )
        try:
            latex, confidence = normalize_ocr_result(engine(_prepare_image(candidate.image)))
            result.latex = latex or None
            result.confidence = confidence
            if not looks_like_formula(latex):
                result.status = "rejected"
                result.reason = "OCR output did not contain a reliable mathematical pattern."
                summary.rejected_count += 1
            else:
                result.status = "recognized"
                summary.recognized_count += 1
        except Exception as exc:  # noqa: BLE001 - third-party boundary.
            result.reason = str(exc)
            summary.failed_count += 1
        summary.results.append(result)

    try:
        _append_formulae(Path(docx_path), summary)
    except FormulaDependencyError as exc:
        summary.warnings.append(str(exc))
    except Exception as exc:  # noqa: BLE001 - third-party boundary.
        summary.warnings.append(f"Recognized formulae could not be added to Word: {exc}")
    if summary.failed_count or summary.rejected_count or summary.warnings:
        summary.status = "completed_with_warnings"
    return summary
